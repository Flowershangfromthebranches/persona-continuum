from __future__ import annotations

import csv
import json
import zipfile
from collections.abc import Sequence
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from persona_continuum.security.paths import validate_zip_members
from persona_continuum.security.validation import SecurityError, ensure_file_size


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    source_type: str
    title: str
    content: str
    metadata: dict[str, object]


SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".jsonl", ".csv", ".html", ".htm", ".docx", ".pdf"}


class SourceLoader:
    max_zip_members = 100
    max_zip_ratio = 100.0

    def __init__(self, max_bytes: int) -> None:
        self.max_bytes = max_bytes

    def load_many(self, paths: Sequence[Path]) -> list[LoadedDocument]:
        loaded: list[LoadedDocument] = []
        for path in paths:
            loaded.extend(self.load(path))
        return loaded

    def load(self, path: Path) -> list[LoadedDocument]:
        if not path.exists():
            raise SecurityError(f"source_not_found:{path}")
        ensure_file_size(path, self.max_bytes)
        suffix = path.suffix.lower()
        if suffix == ".zip":
            return self._load_zip(path)
        if suffix not in SUPPORTED_EXTENSIONS:
            raise SecurityError(f"unsupported_source_type:{suffix}")
        return [self._load_single(path)]

    def _load_zip(self, path: Path) -> list[LoadedDocument]:
        validate_zip_members(path)
        docs: list[LoadedDocument] = []
        with zipfile.ZipFile(path) as archive:
            members = [member for member in archive.infolist() if not member.is_dir()]
            if len(members) > self.max_zip_members:
                raise SecurityError("zip_too_many_members")
            total_uncompressed = sum(member.file_size for member in members)
            if total_uncompressed > self.max_bytes:
                raise SecurityError("zip_total_too_large")
            for member in archive.infolist():
                member_path = Path(member.filename)
                if member.is_dir() or member_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                    continue
                if member.file_size > self.max_bytes:
                    raise SecurityError(f"file_too_large:{member.filename}")
                if (
                    member.compress_size
                    and member.file_size / member.compress_size > self.max_zip_ratio
                ):
                    raise SecurityError("zip_compression_ratio_too_high")
                raw = archive.read(member)
                suffix = member_path.suffix.lower()
                text = self._decode_member(suffix, raw)
                docs.append(
                    LoadedDocument(
                        path=Path(f"{path.name}!{member.filename}"),
                        source_type=suffix.lstrip("."),
                        title=member_path.name,
                        content=self._normalize_text(suffix, text),
                        metadata={
                            "zip": str(path),
                            "member": member.filename,
                            "size": member.file_size,
                        },
                    )
                )
        return docs

    def _load_single(self, path: Path) -> LoadedDocument:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            content = "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs)
        elif suffix == ".pdf":
            reader = PdfReader(str(path))
            content = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            content = path.read_text(encoding="utf-8", errors="replace")
        return LoadedDocument(
            path=path,
            source_type=suffix.lstrip("."),
            title=path.name,
            content=self._normalize_text(suffix, content),
            metadata={"size": path.stat().st_size},
        )

    def _decode_member(self, suffix: str, raw: bytes) -> str:
        if suffix == ".docx":
            return "\n".join(paragraph.text for paragraph in Document(BytesIO(raw)).paragraphs)
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        return raw.decode("utf-8", errors="replace")

    def _normalize_text(self, suffix: str, content: str) -> str:
        if suffix in {".html", ".htm"}:
            return BeautifulSoup(content, "html.parser").get_text("\n")
        if suffix == ".json":
            parsed = json.loads(content)
            return json.dumps(parsed, ensure_ascii=False, indent=2)
        if suffix == ".jsonl":
            rows = [json.loads(line) for line in content.splitlines() if line.strip()]
            return "\n".join(json.dumps(row, ensure_ascii=False) for row in rows)
        if suffix == ".csv":
            rows = list(csv.DictReader(content.splitlines()))
            return "\n".join(json.dumps(row, ensure_ascii=False) for row in rows)
        return content
